"""
Unit tests for AuraDocxConverter.

Run with:  python -m pytest docker/markitdown/aura-plugins/tests/test_docx_converter.py -q
"""

import io
import os
import sys
from datetime import datetime

# Allow running from repo root without installing the package
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from markitdown import StreamInfo

from aura.docx_converter import AuraDocxConverter


def _stream_info(filename: str = "test.docx") -> StreamInfo:
    return StreamInfo(
        extension=".docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )


def _save_docx(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_heading_docx() -> bytes:
    doc = Document()
    doc.core_properties.title = "Spec: Demo"
    doc.core_properties.author = "Aura Test"
    doc.core_properties.created = datetime(2026, 5, 21, 12, 30, 0)

    doc.add_heading("Main Heading", level=1)
    doc.add_paragraph("Introductory paragraph under main heading.")
    doc.add_heading("Heading Level 2", level=2)
    doc.add_paragraph("Detail paragraph under level two.")
    doc.add_heading("Heading Level 3", level=3)
    doc.add_paragraph("Nested paragraph under level three.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"

    return _save_docx(doc)


def _make_no_heading_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Plain paragraph with no heading style.")
    return _save_docx(doc)


def test_accepts_docx_extension():
    c = AuraDocxConverter()
    assert c.accepts(io.BytesIO(b""), _stream_info())


def test_accepts_docx_mime():
    c = AuraDocxConverter()
    si = StreamInfo(
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert c.accepts(io.BytesIO(b""), si)


def test_rejects_xlsx():
    c = AuraDocxConverter()
    si = StreamInfo(
        extension=".xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    assert not c.accepts(io.BytesIO(b""), si)


def test_heading_levels_collapse_to_h3_blocks():
    c = AuraDocxConverter()
    result = c.convert(io.BytesIO(_make_heading_docx()), _stream_info("demo.docx"))
    md = result.markdown
    lines = md.splitlines()

    assert "### Main Heading" in md
    assert "### Heading Level 2" in md
    assert "### Heading Level 3" in md
    assert "## Main Heading" not in lines
    assert "# Main Heading" not in lines


def test_frontmatter_includes_docx_metadata_and_heading_counts():
    c = AuraDocxConverter()
    result = c.convert(io.BytesIO(_make_heading_docx()), _stream_info("demo.docx"))
    md = result.markdown

    assert md.startswith("---")
    assert 'document_title: "Spec: Demo"' in md
    assert 'author: "Aura Test"' in md
    assert 'created: "2026-05-21T12:30:00' in md
    assert "heading_counts:" in md
    assert "  h1: 1" in md
    assert "  h2: 1" in md
    assert "  h3: 1" in md


def test_non_heading_paragraphs_stay_under_current_heading():
    c = AuraDocxConverter()
    result = c.convert(io.BytesIO(_make_heading_docx()), _stream_info("demo.docx"))
    md = result.markdown

    main = md.index("### Main Heading")
    intro = md.index("Introductory paragraph under main heading.")
    h2 = md.index("### Heading Level 2")
    detail = md.index("Detail paragraph under level two.")
    h3 = md.index("### Heading Level 3")
    nested = md.index("Nested paragraph under level three.")

    assert main < intro < h2 < detail < h3 < nested


def test_tables_render_as_markdown_tables():
    c = AuraDocxConverter()
    result = c.convert(io.BytesIO(_make_heading_docx()), _stream_info("demo.docx"))
    md = result.markdown

    assert "| Name | Value |" in md
    assert "| --- | --- |" in md
    assert "| Alpha | 42 |" in md


def test_zero_heading_docx_falls_back_to_builtin_converter():
    c = AuraDocxConverter()
    result = c.convert(io.BytesIO(_make_no_heading_docx()), _stream_info("plain.docx"))
    md = result.markdown

    assert "Plain paragraph with no heading style." in md
    assert "heading_counts:" not in md
    assert not md.startswith("---")
